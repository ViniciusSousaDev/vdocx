import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import threading
from datetime import datetime


# =========================
# EXTRAÇÃO DE DADOS DO PDF
# =========================

def extrair_dados(pdf_path):
    texto_completo = ""
    with pdfplumber.open(pdf_path) as pdf:
        for pagina in pdf.pages:
            texto = pagina.extract_text()
            if texto:
                texto_completo += texto + "\n"

    if not texto_completo.strip():
        raise ValueError(
            "Não foi possível extrair texto do PDF.\n"
            "O arquivo pode ser uma imagem escaneada."
        )

    # DEBUG: descomente para ver o texto extraído
    # print(texto_completo)

    dados = {}

    # --- CNPJ ---
    m = re.search(r"(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})", texto_completo)
    dados["CNPJ"] = m.group(1) if m else "Não encontrado"

    # --- Data de Abertura ---
    m = re.search(r"DATA DE ABERTURA\s*\n\s*(\d{2}/\d{2}/\d{4})", texto_completo, re.IGNORECASE)
    dados["Data de Abertura"] = m.group(1).strip() if m else "Não encontrado"

    # --- Razão Social ---
    m = re.search(r"NOME EMPRESARIAL\s*\n\s*(.+)", texto_completo, re.IGNORECASE)
    dados["Razão Social"] = m.group(1).strip() if m else "Não encontrado"

    # --- Nome Fantasia + Porte ---
    # Cabeçalho: "TÍTULO DO ESTABELECIMENTO (NOME DE FANTASIA)    PORTE"
    # Valor:     "ESTILU CONTABILIDADE                             ME"
    portes = r"(ME|EPP|GRANDE EMPRESA|MÉDIA EMPRESA|MICRO EMPRESA|MICRO-EMPRESA|DEMAIS)"
    m = re.search(
        r"TÍTULO DO ESTABELECIMENTO.*?\n\s*(.+?)\s+" + portes + r"\s*(?:\n|$)",
        texto_completo, re.IGNORECASE
    )
    if m:
        dados["Nome Fantasia"] = m.group(1).strip()
        dados["Porte"] = m.group(2).strip()
    else:
        m2 = re.search(r"TÍTULO DO ESTABELECIMENTO.*?\n\s*(.+)", texto_completo, re.IGNORECASE)
        dados["Nome Fantasia"] = m2.group(1).strip() if m2 else "Não encontrado"
        m3 = re.search(r"PORTE\s*\n\s*" + portes, texto_completo, re.IGNORECASE)
        dados["Porte"] = m3.group(1).strip() if m3 else "Não encontrado"

    # --- CNAE Principal ---
    m = re.search(
        r"CÓDIGO E DESCRIÇÃO DA ATIVIDADE ECONÔMICA PRINCIPAL\s*\n\s*(.+)",
        texto_completo, re.IGNORECASE
    )
    dados["CNAE Principal"] = m.group(1).strip() if m else "Não encontrado"

    # --- CNAEs Secundários ---
    m = re.search(
        r"CÓDIGO E DESCRIÇÃO DAS ATIVIDADES ECONÔMICAS SECUNDÁRIAS\s*\n\s*(.+)",
        texto_completo, re.IGNORECASE
    )
    dados["CNAEs Secundários"] = m.group(1).strip() if m else "Não encontrado"

    # --- Natureza Jurídica ---
    m = re.search(
        r"CÓDIGO E DESCRIÇÃO DA NATUREZA JURÍDICA\s*\n\s*(.+)",
        texto_completo, re.IGNORECASE
    )
    dados["Natureza Jurídica"] = m.group(1).strip() if m else "Não encontrado"

    # --- Logradouro + Número + Complemento ---
    # Cabeçalho: "LOGRADOURO    NÚMERO    COMPLEMENTO"
    # Valor:     "R ORATORIO    1683      SLJ"
    # NÚMERO é sempre numérico → split no primeiro bloco só de dígitos
    m = re.search(
        r"LOGRADOURO.*?NÚMERO.*?COMPLEMENTO.*?\n\s*(.+?)\s+(\d+)\s*(.*?)\s*(?:\n|$)",
        texto_completo, re.IGNORECASE
    )
    if m:
        dados["Logradouro"] = m.group(1).strip()
        dados["Número"] = m.group(2).strip()
        dados["Complemento"] = m.group(3).strip() if m.group(3).strip() else "—"
    else:
        for lbl, chave in [("LOGRADOURO", "Logradouro"), ("NÚMERO", "Número"), ("COMPLEMENTO", "Complemento")]:
            mf = re.search(lbl + r"\s*\n\s*(.+)", texto_completo, re.IGNORECASE)
            dados[chave] = mf.group(1).strip() if mf else "Não encontrado"

    # --- CEP (formato fixo XX.XXX-XXX) ---
    m = re.search(r"(\d{2}\.\d{3}-\d{3})", texto_completo)
    dados["CEP"] = m.group(1) if m else "Não encontrado"

    # --- Bairro, Município, UF ---
    # Linha de valores: "09.280-000   PARQUE DAS NACOES   SANTO ANDRE   SP"
    # UF = 2 letras maiúsculas no final da linha
    m = re.search(
        r"\d{2}\.\d{3}-\d{3}\s+(.+?)\s+([A-Z][A-Z ]+?)\s+([A-Z]{2})\s*(?:\n|$)",
        texto_completo
    )
    if m:
        dados["Bairro"] = m.group(1).strip()
        dados["Município"] = m.group(2).strip()
        dados["UF"] = m.group(3).strip()
    else:
        for lbl, chave in [("BAIRRO/DISTRITO", "Bairro"), ("MUNICÍPIO", "Município"), (r"\bUF\b", "UF")]:
            mf = re.search(lbl + r"\s*\n\s*(.+)", texto_completo, re.IGNORECASE)
            dados[chave] = mf.group(1).strip() if mf else "Não encontrado"

    # --- E-mail (padrão fixo com @) ---
    m = re.search(r"([A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})", texto_completo)
    dados["E-mail"] = m.group(1) if m else "Não encontrado"

    # --- Telefone ---
    m = re.search(r"(\(\d{2}\)\s*\d{4,5}[- ]\d{4})", texto_completo)
    dados["Telefone"] = m.group(1) if m else "Não encontrado"

    # --- Ente Federativo ---
    m = re.search(r"ENTE FEDERATIVO RESPONSÁVEL.*?\n\s*(.+)", texto_completo, re.IGNORECASE)
    dados["Ente Federativo"] = m.group(1).strip() if m else "Não encontrado"

    # --- Situação Cadastral ---
    # Busca a linha de SITUAÇÃO CADASTRAL sem "DATA DA" antes
    m = re.search(
        r"(?:^|(?<!DATA DA\s))SITUAÇÃO CADASTRAL\s*\n\s*(\w[\w ]*?)(?:\n|$)",
        texto_completo, re.IGNORECASE | re.MULTILINE
    )
    dados["Situação Cadastral"] = m.group(1).strip() if m else "Não encontrado"

    # --- Data da Situação Cadastral ---
    m = re.search(
        r"DATA DA SITUAÇÃO CADASTRAL\s*\n\s*(\d{2}/\d{2}/\d{4})",
        texto_completo, re.IGNORECASE
    )
    dados["Data da Situação Cadastral"] = m.group(1).strip() if m else "Não encontrado"

    # --- Capital Social (opcional) ---
    m = re.search(r"CAPITAL SOCIAL.*?\n\s*([\d.,]+)", texto_completo, re.IGNORECASE)
    dados["Capital Social"] = m.group(1).strip() if m else "Não encontrado"

    return dados


# =========================
# GERAÇÃO DO DOCUMENTO WORD
# =========================

ORDEM_SECOES = {
    "Identificação da Empresa": [
        "Razão Social", "Nome Fantasia", "Natureza Jurídica",
        "Capital Social", "Porte", "Data de Abertura", "Ente Federativo",
    ],
    "Situação Cadastral": [
        "Situação Cadastral", "Data da Situação Cadastral",
    ],
    "Endereço": [
        "Logradouro", "Número", "Complemento",
        "Bairro", "CEP", "Município", "UF",
    ],
    "Contato": [
        "E-mail", "Telefone",
    ],
    "Atividade Econômica": [
        "CNAE Principal", "CNAEs Secundários",
    ],
}

CAMPOS_OPCIONAIS = {"Complemento", "E-mail", "CNAEs Secundários", "Ente Federativo", "Capital Social", "Telefone"}


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_para_bg(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def gerar_word(dados, nome_arquivo="empresa.docx"):
    doc = Document()

    secao = doc.sections[0]
    secao.page_width = Inches(8.27)
    secao.page_height = Inches(11.69)
    secao.left_margin = Cm(2)
    secao.right_margin = Cm(2)
    secao.top_margin = Cm(2)
    secao.bottom_margin = Cm(2)

    COR_AZUL_ESCURO = "1F497D"
    COR_AZUL_CLARO = "D6E4F0"

    # Cabeçalho
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMPROVANTE DE INSCRIÇÃO E DE SITUAÇÃO CADASTRAL")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("CNPJ — Receita Federal do Brasil")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()

    # CNPJ em destaque
    p_cnpj = doc.add_paragraph()
    p_cnpj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p_cnpj.add_run("CNPJ: "); rl.bold = True; rl.font.size = Pt(13)
    rv = p_cnpj.add_run(dados.get("CNPJ", "Não encontrado"))
    rv.bold = True; rv.font.size = Pt(13)
    rv.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    # Seções
    for titulo_secao, campos in ORDEM_SECOES.items():
        p_sec = doc.add_paragraph()
        _set_para_bg(p_sec, COR_AZUL_ESCURO)
        rs = p_sec.add_run(f"  {titulo_secao.upper()}  ")
        rs.bold = True; rs.font.size = Pt(10)
        rs.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        tabela = doc.add_table(rows=0, cols=2)
        tabela.style = "Table Grid"

        for campo in campos:
            valor = dados.get(campo, "Não encontrado")
            ausente = valor in ("Não encontrado", "", "*****", "********")
            if ausente and campo in CAMPOS_OPCIONAIS:
                valor = "—"

            linha = tabela.add_row()

            c_lbl = linha.cells[0]; c_lbl.width = Cm(6)
            _set_cell_bg(c_lbl, COR_AZUL_CLARO)
            p_lbl = c_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_before = Pt(3)
            p_lbl.paragraph_format.space_after = Pt(3)
            rl2 = p_lbl.add_run(campo)
            rl2.bold = True; rl2.font.size = Pt(9)
            rl2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            c_val = linha.cells[1]; c_val.width = Cm(11)
            p_val = c_val.paragraphs[0]
            p_val.paragraph_format.space_before = Pt(3)
            p_val.paragraph_format.space_after = Pt(3)
            rv2 = p_val.add_run(valor)
            rv2.font.size = Pt(9)
            if valor == "—":
                rv2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
                rv2.font.italic = True

        doc.add_paragraph()

    # Rodapé
    p_rod = doc.add_paragraph()
    p_rod.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_rod = p_rod.add_run(f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    r_rod.font.size = Pt(7)
    r_rod.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r_rod.font.italic = True

    doc.save(nome_arquivo)
    return nome_arquivo


# =========================
# INTERFACE GRÁFICA (GUI)
# =========================

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Leitor de Cartão CNPJ")
        self.geometry("560x440")
        self.resizable(False, False)
        self.configure(bg="#F0F4F8")
        self.pdf_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.status_text = tk.StringVar(value="Pronto. Selecione o PDF do Cartão CNPJ.")
        self._build_ui()

    def _build_ui(self):
        tk.Label(self, text="Leitor de Cartão CNPJ", font=("Segoe UI", 16, "bold"),
                 bg="#1F497D", fg="white", pady=12).pack(fill=tk.X)
        tk.Label(self, text="Receita Federal do Brasil", font=("Segoe UI", 9),
                 bg="#1F497D", fg="#A8C8E8").pack(fill=tk.X)

        frame = tk.Frame(self, bg="#F0F4F8", padx=24, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(0, weight=1)

        # PDF
        tk.Label(frame, text="Arquivo PDF do Cartão CNPJ:", bg="#F0F4F8",
                 font=("Segoe UI", 10, "bold")).grid(row=0, column=0, sticky="w", pady=(0, 4))
        f_pdf = tk.Frame(frame, bg="#F0F4F8")
        f_pdf.grid(row=1, column=0, sticky="ew", pady=(0, 16))
        f_pdf.columnconfigure(0, weight=1)
        tk.Entry(f_pdf, textvariable=self.pdf_path, font=("Segoe UI", 9),
                 state="readonly", readonlybackground="white", relief=tk.FLAT,
                 highlightthickness=1, highlightbackground="#BCC8D4"
                 ).grid(row=0, column=0, sticky="ew", ipady=6, padx=(0, 8))
        tk.Button(f_pdf, text="Selecionar PDF", command=self._sel_pdf,
                  bg="#1F497D", fg="white", font=("Segoe UI", 9),
                  relief=tk.FLAT, padx=10, cursor="hand2").grid(row=0, column=1)

        # Pasta
        tk.Label(frame, text="Pasta de destino:", bg="#F0F4F8",
                 font=("Segoe UI", 10, "bold")).grid(row=2, column=0, sticky="w", pady=(0, 4))
        f_out = tk.Frame(frame, bg="#F0F4F8")
        f_out.grid(row=3, column=0, sticky="ew", pady=(0, 24))
        f_out.columnconfigure(0, weight=1)
        tk.Entry(f_out, textvariable=self.output_path, font=("Segoe UI", 9),
                 state="readonly", readonlybackground="white", relief=tk.FLAT,
                 highlightthickness=1, highlightbackground="#BCC8D4"
                 ).grid(row=0, column=0, sticky="ew", ipady=6, padx=(0, 8))
        tk.Button(f_out, text="Selecionar Pasta", command=self._sel_pasta,
                  bg="#1F497D", fg="white", font=("Segoe UI", 9),
                  relief=tk.FLAT, padx=10, cursor="hand2").grid(row=0, column=1)

        self.btn = tk.Button(frame, text="⚙  Gerar Documento Word", command=self._processar,
                             bg="#2E7D32", fg="white", font=("Segoe UI", 11, "bold"),
                             relief=tk.FLAT, pady=10, cursor="hand2")
        self.btn.grid(row=4, column=0, sticky="ew")

        self.progress = ttk.Progressbar(frame, mode="indeterminate")
        self.progress.grid(row=5, column=0, sticky="ew", pady=(12, 0))

        tk.Label(frame, textvariable=self.status_text, bg="#F0F4F8",
                 font=("Segoe UI", 9), fg="#555").grid(row=6, column=0, sticky="w", pady=(8, 0))

    def _sel_pdf(self):
        caminho = filedialog.askopenfilename(
            title="Selecionar Cartão CNPJ",
            filetypes=[("Arquivos PDF", "*.pdf"), ("Todos os arquivos", "*.*")]
        )
        if caminho:
            self.pdf_path.set(caminho)
            if not self.output_path.get():
                self.output_path.set(os.path.dirname(caminho))
            self.status_text.set("PDF selecionado. Clique em 'Gerar Documento Word'.")

    def _sel_pasta(self):
        pasta = filedialog.askdirectory(title="Selecionar pasta de destino")
        if pasta:
            self.output_path.set(pasta)

    def _processar(self):
        pdf = self.pdf_path.get()
        pasta = self.output_path.get()
        if not pdf:
            messagebox.showwarning("Atenção", "Selecione o arquivo PDF do Cartão CNPJ.")
            return
        if not pasta:
            messagebox.showwarning("Atenção", "Selecione a pasta de destino.")
            return

        self.btn.config(state=tk.DISABLED)
        self.progress.start(10)
        self.status_text.set("Processando... aguarde.")

        def tarefa():
            try:
                dados = extrair_dados(pdf)
                cnpj_limpo = re.sub(r"[^\d]", "", dados.get("CNPJ", "empresa"))
                nome_arquivo = os.path.join(pasta, f"CNPJ_{cnpj_limpo}.docx")
                gerar_word(dados, nome_arquivo)
                self.after(0, lambda: self._sucesso(nome_arquivo, dados))
            except Exception as e:
                self.after(0, lambda: self._erro(str(e)))

        threading.Thread(target=tarefa, daemon=True).start()

    def _sucesso(self, caminho, dados):
        self.progress.stop()
        self.btn.config(state=tk.NORMAL)
        empresa = dados.get("Razão Social", "")
        self.status_text.set(f"✅ Gerado: {os.path.basename(caminho)}")
        if messagebox.askyesno("Concluído!",
            f"Documento gerado!\n\nEmpresa: {empresa}\nArquivo: {caminho}\n\nAbrir pasta?"):
            os.startfile(os.path.dirname(caminho))

    def _erro(self, mensagem):
        self.progress.stop()
        self.btn.config(state=tk.NORMAL)
        self.status_text.set("❌ Erro ao processar.")
        messagebox.showerror("Erro", f"Ocorreu um erro:\n\n{mensagem}")


if __name__ == "__main__":
    app = App()
    app.mainloop()
